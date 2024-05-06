import mysql.connector
from docx import Document
from docx2pdf import convert

docs = Document()

def list_to_docx(mylist:list()):
    docs.add_heading('Mock Test',0)

    for data in mylist:
        docs.add_paragraph(data,style='List Number')
      
    file = docs.save('Mock_Test.docx')
    return 'Mock_Test.docx'

def docx_to_pdf(myfile):
    outPut_pdf = convert(myfile,'mock_test.pdf')

    return outPut_pdf

conn = mysql.connector.connect(
    host = 'localhost',
    user = 'root',
    password = 'MyNewPass',
    database = 'study'
)

cur = conn.cursor()

cur.execute('select questions from interview_questions order by rand() limit 20')

data = cur.fetchall()

data_list = list()

for x in data:
    # print(list(x))
    # print(type(str(x)))
    data_list.append(x[0])


ques_set = list_to_docx(data_list)
docx_to_pdf(ques_set)
# print(ques_set)
print('pdf file is ready ')